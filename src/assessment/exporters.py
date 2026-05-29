import json
import logging
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

# WeasyPrint Import
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration

logger = logging.getLogger(__name__)

# Suppress noisy logs from pdf generation libraries
# WeasyPrint and FontTools can be very verbose, especially during font subsetting
# The user specifically requested only error logs.
for logger_name in ["weasyprint", "fontTools", "fontTools.subset", "fontTools.ttLib", "pydyf"]:
    logging.getLogger(logger_name).setLevel(logging.ERROR)

RESOURCE_DIR = Path(__file__).parent / "resources" / "fonts"

def get_css_font_faces() -> str:
    """Generates CSS @font-face rules for all available Noto fonts."""
    font_map = {
        "NotoSansDevanagari-Regular.ttf": "NotoSansDevanagari",
        "NotoSansTamil-Regular.ttf": "NotoSansTamil",
        "NotoSansTelugu-Regular.ttf": "NotoSansTelugu",
        "NotoSansKannada-Regular.ttf": "NotoSansKannada",
        "NotoSansMalayalam-Regular.ttf": "NotoSansMalayalam",
        "NotoSansBengali-Regular.ttf": "NotoSansBengali",
        "NotoSansGujarati-Regular.ttf": "NotoSansGujarati",
        "NotoSansGurmukhi-Regular.ttf": "NotoSansGurmukhi"
    }
    
    css = []
    for filename, font_family in font_map.items():
        font_path = RESOURCE_DIR / filename
        if font_path.exists():
            # WeasyPrint needs file:// URI for local files or absolute paths
            css.append(f"""
            @font-face {{
                font-family: '{font_family}';
                src: url('file://{font_path.absolute()}');
            }}""")
            
    return "\n".join(css)

def generate_html_content(assessment_data: dict) -> str:
    """Constructs the HTML report string."""
    blueprint = assessment_data.get("blueprint", {})
    questions_obj = assessment_data.get("questions", {})
    
    # Audit Data
    prompt_ver = blueprint.get("prompt_version", "N/A")
    api_ver = blueprint.get("api_version", "N/A")
    scope = blueprint.get('assessment_scope_summary', 'N/A')
    
    # CSS
    font_faces = get_css_font_faces()
    # Pango Stack: Put specific fonts first, then sans-serif fallback
    font_stack = "'NotoSansMalayalam', 'NotoSansDevanagari', 'NotoSansTamil', 'NotoSansTelugu', 'NotoSansKannada', 'NotoSansBengali', 'NotoSansGujarati', 'NotoSansGurmukhi', sans-serif"

    html_parts = ["""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            %s
            body {
                font-family: %s;
                font-size: 11pt;
                line-height: 1.5;
                color: #333;
                margin: 40px;
            }
            h1 { font-size: 24pt; color: #2c3e50; border-bottom: 2px solid #eee; padding-bottom: 10px; }
            h2 { font-size: 18pt; color: #34495e; margin-top: 30px; }
            h3 { font-size: 14pt; color: #7f8c8d; }
            .audit-table { width: 100%%; border-collapse: collapse; margin-bottom: 20px; }
            .audit-table th, .audit-table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            .audit-table th { background-color: #f2f2f2; }
            .question-block { margin-bottom: 25px; page-break-inside: avoid; }
            .question-text { font-weight: bold; font-size: 12pt; margin-bottom: 8px; }
            .options-list { list-style-type: none; padding-left: 20px; }
            .options-list li { margin-bottom: 4px; }
            .reasoning-box { background-color: #f8f9fa; border-left: 4px solid #3498db; padding: 10px; margin-top: 10px; font-size: 10pt; }
            .correct { color: #27ae60; font-weight: bold; }
        </style>
    </head>
    <body>
        <h1>Course Assessment Report</h1>
        
        <p><b>Assessment Scope:</b> %s</p>
        
        <h3>Audit Information</h3>
        <table class="audit-table">
            <tr><th>Field</th><th>Value</th></tr>
            <tr><td>Prompt Version</td><td>%s</td></tr>
            <tr><td>API Version</td><td>%s</td></tr>
        </table>
        
        <h2>Questions & Reasoning</h2>
    """ % (font_faces, font_stack, scope, prompt_ver, api_ver)]

    # Dynamic Questions
    q_counter = 1
    for q_type, q_list in questions_obj.items():
        html_parts.append(f"<h3>{q_type} ({len(q_list)})</h3>")
        
        for q in q_list:
            if q_type == "MTF Question":
                default_q_txt = q.get("matching_context", "Match the following items appropriately:")
            else:
                default_q_txt = "N/A"
            
            q_txt = q.get("question_text", default_q_txt)
            q_html = f"""
            <div class="question-block">
                <div class="question-text">Q{q_counter}: {q_txt}</div>
                <div style="font-size: 0.85em; color: #666; margin-bottom: 8px;"><i>Source Course: {q.get('course_name', 'N/A')}</i></div>
            """
            
            # Options / Body
            ALPHA = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if q_type == "Multiple Choice Question":
                options = q.get("options", [])
                opts_html = "".join([f"<li>{ALPHA[i]}. {o.get('text', '')}</li>" for i, o in enumerate(options)])
                correct_idx = q.get('correct_option_index')
                correct_set = {int(correct_idx)} if correct_idx is not None else set()
                correct_labels = [
                    ALPHA[i] for i, o in enumerate(options)
                    if (int(o["index"]) if o.get("index") is not None else i) in correct_set
                ]
                q_html += f"<ul class='options-list'>{opts_html}</ul>"
                q_html += f"<div class='correct'>Correct Answer: {correct_labels[0] if correct_labels else 'N/A'}</div>"

            elif q_type == "MTF Question":
                pairs_html = "".join([f"<li>{ALPHA[i]}. {p.get('left')} &rarr; {p.get('right')}</li>" for i, p in enumerate(q.get("pairs", []))])
                q_html += f"<ul class='options-list'>{pairs_html}</ul>"

            elif q_type == "Multi-Choice Question":
                options = q.get("options", [])
                opts_html = "".join([f"<li>{ALPHA[i]}. {o.get('text', '')}</li>" for i, o in enumerate(options)])
                corr = q.get('correct_option_index')
                correct_set = {int(x) for x in corr} if isinstance(corr, list) else ({int(corr)} if corr is not None else set())
                correct_labels = [
                    ALPHA[i] for i, o in enumerate(options)
                    if (int(o["index"]) if o.get("index") is not None else i) in correct_set
                ]
                q_html += f"<ul class='options-list'>{opts_html}</ul>"
                q_html += f"<div class='correct'>Correct Options: {', '.join(correct_labels) if correct_labels else 'N/A'}</div>"

            elif q_type == "True/False Question":
                q_html += "<ul class='options-list'><li>A. True</li><li>B. False</li></ul>"
                q_html += f"<div class='correct'>Correct Answer: {q.get('correct_answer')}</div>"
            
            else:
                q_html += f"<div class='correct'>Answer: {q.get('correct_answer')}</div>"

            # Reasoning Box
            rs = q.get("reasoning", {})
            ar = q.get("answer_rationale", {})
            kcm = rs.get("competency_alignment", {}).get("kcm", {})
            q_html += f"""
                <div class="reasoning-box">
                    <b>Rationale:</b> {ar.get('correct_answer_explanation', 'N/A')}<br/>
                    <b>Bloom's Level:</b> {q.get('blooms_level', 'N/A')} ({rs.get('blooms_level_justification', 'N/A')})<br/>
                    <b>Learning Objective:</b> {rs.get('learning_objective_alignment', 'N/A')}<br/>
                    <b>Competency:</b> {kcm.get('competency_area', 'N/A')} - {kcm.get('competency_theme', 'N/A')}<br/>
                    <b>Relevance:</b> {q.get('relevance_percentage', 'N/A')}%
                </div>
            </div>
            """
            html_parts.append(q_html)
            q_counter += 1

    html_parts.append("</body></html>")
    return "\n".join(html_parts)

def generate_pdf(assessment_data: dict, output_path: Path):
    """Generates a PDF report using WeasyPrint (HTML-to-PDF)."""
    try:
        html_content = generate_html_content(assessment_data)
        font_config = FontConfiguration()
        HTML(string=html_content).write_pdf(target=str(output_path), font_config=font_config)
        logger.info(f"Generated PDF with WeasyPrint: {output_path}")
    except Exception as e:
        logger.error(f"WeasyPrint PDF Generation Failed: {e}")
        # Fallback? No, fail explicitly is better than bad boxes.
        raise


def generate_docx(assessment_data: dict, output_path: Path):
    """Generates a DOCX report from the assessment JSON data."""
    doc = Document()
    doc.add_heading('Course Assessment Report', 0)

    blueprint = assessment_data.get("blueprint", {})
    doc.add_paragraph(f"Assessment Scope: {blueprint.get('assessment_scope_summary', 'N/A')}")
    
    # Audit Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Prompt Version'
    row_cells[1].text = str(blueprint.get("prompt_version", "N/A"))
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'API Version'
    row_cells[1].text = str(blueprint.get("api_version", "N/A"))
    
    doc.add_heading('Questions & Reasoning', level=1)
    
    questions_obj = assessment_data.get("questions", {})
    
    for q_type, q_list in questions_obj.items():
        doc.add_heading(f"{q_type} ({len(q_list)})", level=2)
        
        for i, q in enumerate(q_list, 1):
            if q_type == "MTF Question":
                default_q_txt = q.get("matching_context", "Match the following items appropriately:")
            else:
                default_q_txt = "N/A"
                
            q_para = doc.add_paragraph()
            q_para.add_run(f"{i}. {q.get('question_text', default_q_txt)}").bold = True

            # Course Tag
            c_para = doc.add_paragraph()
            c_para.add_run(f"Source Course: {q.get('course_name', 'N/A')}").italic = True

            ALPHA = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if q_type == "Multiple Choice Question":
                options = q.get("options", [])
                for i, opt in enumerate(options):
                    doc.add_paragraph(f"{ALPHA[i]}. {opt.get('text', '')}", style='List Bullet')
                correct_idx = q.get('correct_option_index')
                correct_set = {int(correct_idx)} if correct_idx is not None else set()
                correct_labels = [
                    ALPHA[i] for i, o in enumerate(options)
                    if (int(o["index"]) if o.get("index") is not None else i) in correct_set
                ]
                p = doc.add_paragraph()
                p.add_run(f"Correct Answer: {correct_labels[0] if correct_labels else 'N/A'}").bold = True

            elif q_type == "MTF Question":
                for i, p_item in enumerate(q.get("pairs", [])):
                    doc.add_paragraph(f"{ALPHA[i]}. {p_item.get('left')} -> {p_item.get('right')}", style='List Bullet')

            elif q_type == "Multi-Choice Question":
                options = q.get("options", [])
                for i, opt in enumerate(options):
                    doc.add_paragraph(f"{ALPHA[i]}. {opt.get('text', '')}", style='List Bullet')
                corr = q.get('correct_option_index')
                correct_set = {int(x) for x in corr} if isinstance(corr, list) else ({int(corr)} if corr is not None else set())
                correct_labels = [
                    ALPHA[i] for i, o in enumerate(options)
                    if (int(o["index"]) if o.get("index") is not None else i) in correct_set
                ]
                p = doc.add_paragraph()
                p.add_run(f"Correct Options: {', '.join(correct_labels) if correct_labels else 'N/A'}").bold = True

            elif q_type == "True/False Question":
                 doc.add_paragraph("A. True", style='List Bullet')
                 doc.add_paragraph("B. False", style='List Bullet')
                 p = doc.add_paragraph()
                 p.add_run(f"Correct Answer: {q.get('correct_answer')}").bold = True
            
            else:
                 p = doc.add_paragraph()
                 p.add_run(f"Answer: {q.get('correct_answer')}").bold = True

            # Reasoning
            reasoning = q.get("reasoning", {})
            ar = q.get("answer_rationale", {})
            kcm = reasoning.get("competency_alignment", {}).get("kcm", {})
            
            r_para = doc.add_paragraph()
            r_para.add_run(f"Rationale: {ar.get('correct_answer_explanation', 'N/A')}").italic = True

            blooms_para = doc.add_paragraph()
            blooms_para.add_run(f"Bloom's: {q.get('blooms_level', 'N/A')} ({reasoning.get('blooms_level_justification', 'N/A')}) | Relevance: {q.get('relevance_percentage', 'N/A')}%")

            lo_para = doc.add_paragraph()
            lo_para.add_run(f"Learning Objective: {reasoning.get('learning_objective_alignment', 'N/A')}")

            comp_para = doc.add_paragraph()
            comp_para.add_run(f"Competency: {kcm.get('competency_area', 'N/A')} - {kcm.get('competency_theme', 'N/A')}")

            doc.add_paragraph()
            
    doc.save(str(output_path))
